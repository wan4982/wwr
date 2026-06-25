from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


INPUT = Path(r"D:\StudyProjects\campus_assistant\docs\202431243 王婉如 实训报告-第1天已填.docx")
OUTPUT = Path(r"D:\StudyProjects\campus_assistant\docs\202431243 王婉如 实训报告-每日记录已填.docx")


def set_cell_text(cell, text: str, font_size: float = 10.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.line_spacing = 1.2
    paragraph.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(text.split("\n")):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(font_size)


DAY_CONTENT = {
    1: {
        "date": "2026年6月23日",
        "place": "机房 / 宿舍",
        "task": "完成向量库与检索模块开发：使用TF-IDF方式构建本地轻量向量索引，实现关键词扩展、相似度计算、Top-K检索，并生成 vector_db/tfidf_index.json。",
        "problems": [
            ("实训指导书建议使用Chroma和嵌入模型，但本地安装和模型下载受网络环境影响较大。", "采用TF-IDF实现轻量级本地向量检索，保证没有外部模型时项目也能运行演示。"),
            ("中文短问题容易因为词面不一致导致检索偏差，例如“发烧”与“病假”不是完全相同的词。", "增加查询扩展词典，把发烧、生病、病假、请假等相关词关联起来，提高检索命中率。"),
            ("奖学金、请假等类别中有相似问题，单纯计算全文相似度可能排序不准。", "在检索得分中增加问题字段匹配加权，让用户问题和知识库question更接近的记录优先返回。"),
            ("CSV中每条问答长度较短，不确定是否需要继续切分。", "分析后认为每条问答已经是较小知识片段，暂时不再切分；后续若加入长制度原文再进行chunk切分。"),
            ("检索结果需要便于调试和报告展示。", "在检索函数中返回记录和相似度，并在Streamlit中增加“检索过程”页面展示Top-K结果。"),
        ],
        "code": """# src/retriever.py：本地TF-IDF检索核心
def search(self, query, top_k=DEFAULT_TOP_K, category=None):
    query_vector = self._vectorize(tokenize(expand_query(query)))
    query_tokens = set(tokenize(query))
    scored = []
    for record, doc_vector in zip(self.records, self.doc_vectors):
        if category and record.category != category:
            continue
        score = self._cosine(query_vector, doc_vector)
        question_tokens = set(tokenize(record.question))
        if query_tokens and question_tokens:
            score += 0.35 * (len(query_tokens & question_tokens) / len(query_tokens))
        if query.strip() == record.question.strip():
            score += 1.0
        scored.append((record, score))
    scored.sort(key=lambda item: item[1], reverse=True)
    return [item for item in scored[:top_k] if item[1] > 0]""",
        "summary": "第二天主要完成了知识库检索模块。由于本地环境不一定能顺利安装Chroma和下载中文嵌入模型，我选择先实现一个轻量级TF-IDF向量检索方案。这个方案虽然不如专业嵌入模型语义能力强，但优点是稳定、离线可运行，适合实训演示。我还针对校园问答场景增加了查询扩展和问题字段加权，使“病假、请假、发烧”等相关表达能够更好地命中。通过这一天的开发，我理解到向量库并不是简单存数据，而是要考虑文本切分、字段设计、相似度排序和可解释的检索结果。",
    },
    2: {
        "date": "2026年6月24日",
        "place": "机房 / 宿舍",
        "task": "完成RAG问答模块：设计校园助手提示词，组合检索结果生成回答；在未配置大模型API时使用本地抽取式回答，配置DeepSeek API后可调用大模型增强。",
        "problems": [
            ("直接返回最相似的一条记录时，回答有时不够完整。", "对请假类问题增加流程型组合回答，把病假、事假、销假、特殊情况等多条规则组织成完整答案。"),
            ("如果问题不在知识库中，模型可能编造答案。", "在提示词和本地回答中加入约束：没有相关规则时回答“我不确定，建议咨询辅导员”。"),
            ("不同问题需要不同回答形式，例如普通问答和流程步骤不一样。", "在RAG模块中判断泛请假、病假、事假、销假等问题，分别生成汇总型或流程型回答。"),
            ("大模型API可能没有Key或网络不稳定。", "保留DeepSeek调用入口，同时提供本地回答兜底，保证项目不依赖外部API也能演示。"),
            ("回答需要有依据，方便老师检查数据来源。", "每次回答末尾输出参考来源，来源来自CSV中的source字段。"),
        ],
        "code": """# src/rag.py：RAG问答入口
def rag_answer(question, retriever=None, use_llm=True):
    active_retriever = retriever or build_retriever()

    specific_leave = _specific_leave_answer(question, active_retriever)
    if specific_leave:
        answer, results = specific_leave
        return {"answer": answer, "context": _format_context(results), "results": results}

    if _is_general_leave_question(question):
        answer, results = _leave_summary(active_retriever)
        return {"answer": answer, "context": _format_context(results), "results": results}

    results = active_retriever.search(question, top_k=3)
    context = _format_context(results)
    answer = _deepseek_answer(question, context) if use_llm and results else None
    if not answer:
        answer = _local_answer(question, results)
    return {"answer": answer, "context": context, "results": results}""",
        "summary": "第三天主要实现了RAG问答功能。RAG的核心流程是先检索知识库，再把相关规则作为上下文生成回答。我在项目中实现了两种回答方式：如果配置了DeepSeek API，可以调用大模型生成更自然的回答；如果没有API Key，则使用本地抽取式回答，保证项目仍然可运行。调试过程中我发现，请假这类问题不能只返回一条记录，否则会漏掉申请入口、材料、审批和销假等步骤，因此我专门写了请假流程组合回答。通过这一天，我认识到RAG系统不仅要能检索，还要根据问题类型组织答案。",
    },
    3: {
        "date": "2026年6月25日",
        "place": "机房 / 宿舍",
        "task": "完成智能体模块：实现意图识别，支持校历周次查询、绩点计算工具调用，并将普通校园生活问题转入RAG问答流程。",
        "problems": [
            ("所有问题都走RAG会导致计算类问题不方便处理。", "增加智能体意图识别，对“现在第几周”和“绩点计算”分别调用工具函数。"),
            ("奖学金问题中也包含“绩点”二字，容易被误判成绩点计算。", "收窄GPA工具触发条件，只有明确出现“绩点计算”或包含具体分数时才调用计算工具。"),
            ("周次计算需要一个学期开始日期。", "在配置文件中设置学期开始日期，再用当前日期和开始日期相差天数计算教学周。"),
            ("绩点计算需要处理不同输入格式。", "使用正则提取用户输入中的分数，支持“85,90,78”等格式，并检查分数范围。"),
            ("多轮对话需要保留最近历史。", "在智能体类中维护conversation_history列表，保留最近10条消息作为基础对话记忆。"),
        ],
        "code": """# src/agent.py：智能体路由
def chat(self, user_input, use_llm=True):
    question = user_input.strip()
    self.conversation_history.append({"role": "user", "content": question})

    if self._is_week_query(question):
        response = get_current_week()
    elif self._is_gpa_query(question):
        response = calculate_gpa(question)
    else:
        response = rag_answer(question, self.retriever, use_llm=use_llm)["answer"]

    self.conversation_history.append({"role": "assistant", "content": response})
    self.conversation_history = self.conversation_history[-10:]
    return response""",
        "summary": "第四天主要完成智能体部分。相比单纯的问答系统，智能体需要先判断用户意图，再决定调用工具还是进入RAG检索。我实现了两个常用工具：校历周次查询和绩点计算。周次工具根据学期开始日期自动计算当前是第几周；绩点工具可以从用户输入中提取分数并计算平均绩点。开发中比较容易出错的是意图识别，例如“奖学金要多少绩点”不应该进入GPA计算，所以我调整了判断规则。通过这一天，我理解到智能体的关键是把不同能力组织起来，而不是所有问题都交给大模型。",
    },
    4: {
        "date": "2026年6月26日",
        "place": "机房 / 宿舍",
        "task": "完成Streamlit Web界面和本地部署测试：实现智能问答、检索过程展示、知识库浏览、快捷问题和DeepSeek API开关，并在本地启动服务进行验证。",
        "problems": [
            ("Streamlit首次启动会写入用户目录配置，受权限限制可能失败。", "将HOME和USERPROFILE设置到项目目录，避免写入受限目录。"),
            ("修改CSV后页面仍可能使用旧缓存。", "给缓存函数加入数据版本号，版本由CSV修改时间和文件大小组成，数据变化后自动重载。"),
            ("页面示例问题如果与CSV不一致，会给人造成答非所问的感觉。", "把快捷示例调整为知识库中真实存在的问题，如宿舍熄灯、请假、奖学金、报修、一卡通等。"),
            ("调试时需要确认服务是否真正可访问。", "使用Invoke-WebRequest检查本地端口返回200，并用命令行测试核心问答逻辑。"),
            ("报告最终需要截图，但当前阶段先保证功能可用。", "保留智能问答、检索过程和知识库三个页面，后续完成报告时再统一截图整理。"),
        ],
        "code": """# src/app.py：Streamlit缓存随CSV变化自动刷新
def data_version():
    stat = DATA_FILE.stat()
    return f"{stat.st_mtime_ns}-{stat.st_size}"

@st.cache_resource
def get_agent(version):
    return CampusAgent(build_retriever())

@st.cache_data
def load_campus_data(version):
    return pd.read_csv(DATA_FILE)

version = data_version()
agent = get_agent(version)""",
        "summary": "第五天主要完成项目集成和本地部署。前几天已经完成数据、检索、RAG和智能体，这一天重点是把功能放到可操作的Web界面中。我使用Streamlit实现了三个页面：智能问答用于和助手对话，检索过程用于查看Top-K结果和相似度，知识库页面用于浏览CSV数据。部署测试时遇到缓存和权限问题，例如CSV更新后页面仍可能使用旧数据，后来通过数据版本号解决。最终项目可以在本地浏览器访问，并能回答请假、奖学金、报修、一卡通、选课、宿舍管理等问题。通过这一天，我体会到AI应用不仅要算法能跑，还要界面清楚、部署稳定、测试可复现。",
    },
}


def fill_table(table, content: dict[str, object]) -> None:
    set_cell_text(table.cell(0, 1), content["date"])
    set_cell_text(table.cell(0, 3), content["place"])
    set_cell_text(table.cell(1, 1), content["task"])

    set_cell_text(table.cell(2, 0), "一、遇到的问题及其解决方法")
    set_cell_text(table.cell(3, 0), "问题描述")
    set_cell_text(table.cell(3, 2), "解决方法")
    for row_index, (problem, solution) in enumerate(content["problems"], start=4):
        set_cell_text(table.cell(row_index, 0), problem)
        set_cell_text(table.cell(row_index, 2), solution)

    set_cell_text(table.cell(9, 0), "二、关键实现代码")
    set_cell_text(table.cell(10, 0), content["code"], font_size=9)
    if len(table.rows) >= 13:
        set_cell_text(table.cell(11, 0), "三、实训小结（不少于100字）")
        set_cell_text(table.cell(12, 0), content["summary"])
    else:
        set_cell_text(
            table.cell(11, 0),
            f"三、实训小结（不少于100字）\n{content['summary']}",
        )


def main() -> None:
    doc = Document(INPUT)
    for table_index, content in DAY_CONTENT.items():
        fill_table(doc.tables[table_index], content)
    doc.save(OUTPUT)
    print(f"已填写第2-5天每日记录：{OUTPUT}")


if __name__ == "__main__":
    main()
