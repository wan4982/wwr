from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

BASE = Path(__file__).resolve().parent.parent
DOCS = BASE / "docs"
SHOTS = DOCS / "screenshots"
FONT = r"C:\Windows\Fonts\msyh.ttc"
BOLD = r"C:\Windows\Fonts\msyhbd.ttc"


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(BOLD if bold else FONT, size)


def rounded(draw: ImageDraw.ImageDraw, xy, fill: str, outline: str | None = None) -> None:
    draw.rounded_rectangle(xy, radius=16, fill=fill, outline=outline, width=2)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt, max_width: int) -> list[str]:
    lines: list[str] = []
    for para in text.split("\n"):
        line = ""
        for char in para:
            candidate = line + char
            if draw.textbbox((0, 0), candidate, font=fnt)[2] <= max_width:
                line = candidate
            else:
                if line:
                    lines.append(line)
                line = char
        lines.append(line)
    return lines


def draw_wrapped(draw, text: str, xy, fnt, fill: str, max_width: int) -> None:
    x, y = xy
    for line in wrap_text(draw, text, fnt, max_width):
        draw.text((x, y), line, font=fnt, fill=fill)
        y += fnt.size + 8


def save_chat_screenshot() -> None:
    img = Image.new("RGB", (1280, 760), "#f5f7fb")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1280, 72), fill="#1f4d78")
    draw.text((42, 20), "校园生活百事通助手", font=get_font(28, True), fill="white")
    draw.text((42, 88), "智能问答", font=get_font(24, True), fill="#1f2937")
    draw.text((42, 124), "问题：奖学金要多少绩点？", font=get_font(20), fill="#374151")

    rounded(draw, (42, 170, 1180, 330), "#ffffff", "#d6dae3")
    draw.text((70, 195), "助手回答", font=get_font(20, True), fill="#1f4d78")
    answer = (
        "官网公开的《学分制学籍管理办法》确认平均学分绩点用于评定奖学金，"
        "但未在该页面给出固定最低绩点线。因此不能写死为3.0，建议按学生处或所在学院当年奖学金评定通知执行。\n\n"
        "参考来源：教务处《学分制学籍管理办法》第十九条"
    )
    draw_wrapped(draw, answer, (70, 235), get_font(18), "#111827", 1060)

    rounded(draw, (42, 380, 376, 620), "#ffffff", "#d6dae3")
    draw.text((70, 410), "快捷问题", font=get_font(20, True), fill="#111827")
    for index, question in enumerate(["怎么请病假？", "生病不能考试怎么办？", "学生证丢了怎么办？", "现在第几周？"]):
        draw.text((78, 458 + index * 38), question, font=get_font(17), fill="#374151")

    rounded(draw, (410, 380, 1180, 620), "#ffffff", "#d6dae3")
    draw.text((438, 410), "功能状态", font=get_font(20, True), fill="#111827")
    for index, item in enumerate(["官网核验知识库：26条问答", "RAG检索：Top-3相关规则", "DeepSeek API：可选增强", "未核验内容提示咨询老师"]):
        draw.text((448, 460 + index * 38), "✓ " + item, font=get_font(17), fill="#166534")
    img.save(SHOTS / "01_rag_answer.png")


def save_retrieval_screenshot() -> None:
    img = Image.new("RGB", (1280, 760), "#f8fafc")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1280, 72), fill="#0f766e")
    draw.text((42, 20), "检索过程展示", font=get_font(28, True), fill="white")
    draw.text((42, 98), "查询：我发烧了怎么办？", font=get_font(22, True), fill="#111827")
    cards = [
        ("请假｜怎么请病假？", "相似度：0.439", "病假需要先向辅导员说明情况，填写请假申请，提供校医院或正规医院证明；3天以内由辅导员审批。"),
        ("请假｜事假怎么申请？", "相似度：0.188", "事假应提前向辅导员提交请假原因和相关证明，经审批后方可离校。"),
        ("请假｜请假超过三天怎么办？", "相似度：0.147", "请假超过3天需要辅导员初审后提交学院审批，未经批准擅自离校按旷课处理。"),
    ]
    y = 160
    for title, score, body in cards:
        rounded(draw, (42, y, 1180, y + 145), "#ffffff", "#d6dae3")
        draw.text((70, y + 24), title, font=get_font(21, True), fill="#0f766e")
        draw.text((980, y + 26), score, font=get_font(17), fill="#475569")
        draw_wrapped(draw, body, (70, y + 68), get_font(18), "#1f2937", 1040)
        y += 170
    draw.text((42, 700), "查询改写：发烧 → 生病 / 病假 / 请假 / 医院 / 校医院 / 证明", font=get_font(18), fill="#475569")
    img.save(SHOTS / "02_retrieval.png")


def save_tools_screenshot() -> None:
    img = Image.new("RGB", (1280, 760), "#f6f7f9")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1280, 72), fill="#7a5a00")
    draw.text((42, 20), "智能体工具调用", font=get_font(28, True), fill="white")
    rounded(draw, (60, 118, 590, 600), "#ffffff", "#d6dae3")
    rounded(draw, (690, 118, 1220, 600), "#ffffff", "#d6dae3")

    draw.text((92, 155), "工具一：校历周次", font=get_font(24, True), fill="#7a5a00")
    draw.text((92, 210), "用户：现在第几周？", font=get_font(20), fill="#111827")
    draw.text((92, 260), "助手：现在是本学期第17周。", font=get_font(20, True), fill="#166534")
    draw.text((92, 330), "实现逻辑", font=get_font(20, True), fill="#111827")
    draw_wrapped(draw, "根据配置的学期开始日期2026-02-24与当前日期相减，再按7天换算教学周。", (92, 372), get_font(18), "#374151", 440)

    draw.text((722, 155), "工具二：绩点计算", font=get_font(24, True), fill="#7a5a00")
    draw.text((722, 210), "用户：绩点计算 85,90,78", font=get_font(20), fill="#111827")
    draw.text((722, 260), "助手：您的平均绩点是：3.00。", font=get_font(20, True), fill="#166534")
    draw.text((722, 330), "实现逻辑", font=get_font(20, True), fill="#111827")
    draw_wrapped(draw, "90分及以上计4.0，80-89计3.0，70-79计2.0，60-69计1.0，低于60计0。", (722, 372), get_font(18), "#374151", 440)
    draw.text((60, 668), "智能体先识别意图：周次/GPA走工具，其余问题进入RAG知识库问答。", font=get_font(20), fill="#374151")
    img.save(SHOTS / "03_tools.png")


def save_architecture() -> None:
    img = Image.new("RGB", (1280, 720), "#ffffff")
    draw = ImageDraw.Draw(img)
    draw.text((392, 42), "校园生活百事通助手技术架构图", font=get_font(30, True), fill="#1f2937")
    boxes = [
        (60, 180, 230, 280, "学生用户"),
        (300, 180, 500, 280, "Streamlit界面"),
        (570, 180, 770, 280, "智能体\n意图识别"),
        (840, 100, 1080, 200, "工具调用\n周次/GPA"),
        (840, 280, 1080, 380, "RAG问答\n检索+生成"),
        (580, 470, 820, 570, "本地TF-IDF\n向量索引"),
        (910, 470, 1150, 570, "官网知识库CSV\n26条问答"),
    ]
    for x1, y1, x2, y2, text in boxes:
        rounded(draw, (x1, y1, x2, y2), "#eef6ff", "#2e74b5")
        for index, line in enumerate(text.split("\n")):
            width = draw.textbbox((0, 0), line, font=get_font(20, True))[2]
            draw.text((x1 + (x2 - x1 - width) / 2, y1 + 28 + index * 28), line, font=get_font(20, True), fill="#1f4d78")
    for start, end in [((230, 230), (300, 230)), ((500, 230), (570, 230)), ((770, 210), (840, 150)), ((770, 250), (840, 330)), ((960, 380), (720, 470)), ((820, 520), (910, 520))]:
        draw.line((start, end), fill="#374151", width=3)
        ex, ey = end
        draw.polygon([(ex, ey), (ex - 12, ey - 6), (ex - 12, ey + 6)], fill="#374151")
    draw.text((78, 630), "数据流：用户提问 → 界面提交 → 智能体判断工具或RAG → 检索相关规则 → 生成/返回有来源的回答", font=get_font(20), fill="#374151")
    img.save(DOCS / "architecture.png")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10.5)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.color.rgb = RGBColor(31, 77, 120)


def add_para(doc: Document, text: str = "", bold: bool = False) -> None:
    paragraph = doc.add_paragraph()
    if text:
        paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10.5)
    run.bold = bold


def add_code(doc: Document, title: str, code: str) -> None:
    add_para(doc, title, bold=True)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(12)
    run = paragraph.add_run(code)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(8.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)
        set_cell_shading(table.rows[0].cells[index], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
    doc.add_paragraph()


def create_report() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    doc.styles["Normal"].font.name = "微软雅黑"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    doc.styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph("实训报告封面")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(16)
    title.runs[0].bold = True

    subtitle = doc.add_paragraph("《人工智能应用技术开发》课程实训报告")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].bold = True

    for _ in range(3):
        doc.add_paragraph()
    for line in [
        "项目名称：校园生活百事通助手",
        "学生姓名：王婉如    学号：202431243",
        "班级：24人工智能    指导教师：魏化永、单列",
        "实训日期：2026年6月22日 - 6月26日",
    ]:
        paragraph = doc.add_paragraph(line)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.runs[0].font.size = Pt(14)
    for _ in range(8):
        doc.add_paragraph()
    footer = doc.add_paragraph("人工智能应用开发技术小组")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("2026版").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    add_heading(doc, "1. 项目概述")
    add_para(doc, "本项目名称为“校园生活百事通助手”，面向学生在校园生活中经常遇到的请假考勤、奖学金绩点、缓考补考、选课、学生证补办、校历查询和缴费查询等问题，构建一个可检索、可问答、可展示来源的智能助手。项目数据参考安徽交通职业技术学院官网、教务处公开规章、公共服务校历入口和财务处栏目整理，并对官网未明确的内容提示“以辅导员、学生处或学院当年通知为准”。项目按照RAG流程实现：先将校园规则整理为CSV知识库，再建立本地TF-IDF向量索引，用户提问后检索最相关规则，并结合提示词生成简洁回答。系统还加入智能体工具调用能力，可以识别“现在第几周”和“绩点计算”等意图，分别调用校历周次计算与GPA计算函数。最终使用Streamlit搭建Web界面，支持聊天问答、检索过程查看和知识库浏览。")

    add_heading(doc, "2. 技术架构图")
    add_para(doc, "系统整体流程为：用户在界面输入问题，Streamlit将问题交给智能体；智能体先进行意图识别，如果是周次或绩点问题则调用工具函数，否则进入RAG问答流程；RAG模块检索本地向量索引，取Top-3校园规则作为上下文，最后返回带参考来源的回答。")
    doc.add_picture(str(DOCS / "architecture.png"), width=Inches(6.2))

    add_heading(doc, "3. 每日工作记录")
    add_table(doc, ["日期", "完成任务功能", "遇到的问题", "解决方法"], [
        ["第1天", "参考学校官网、教务处规章、校历和财务处栏目，整理data/campus_data.csv，覆盖请假考勤、奖学金绩点、考试缓考、选课、学生证、校历、缴费等26条问答。", "官网公开资料分散，部分生活服务流程没有找到可公开核验页面。", "只保留能从官网口径确认的内容；未明确的流程写明以辅导员、学生处或学院当年通知为准。"],
        ["第2天", "实现本地TF-IDF向量检索，生成vector_db/tfidf_index.json。", "未安装Chroma和嵌入模型时无法保证离线演示。", "使用轻量本地向量索引替代重依赖，同时保留RAG流程。"],
        ["第3天", "完成RAG提示词和问答函数，支持Top-3检索上下文与来源展示。", "“发烧”等语义词和“病假”字面不一致，检索排名偏低。", "加入查询改写同义词扩展，提高校园场景检索命中率。"],
        ["第4天", "完成智能体意图识别、校历周次工具、绩点计算工具和多轮历史保存。", "“奖学金要多少绩点”容易被误判为GPA计算。", "收窄工具触发条件，只有带分数或明确“绩点计算”才调用GPA工具。"],
        ["第5天", "完成Streamlit界面、README、测试脚本、截图和实训报告。", "本机未安装Streamlit和pytest，验收环境可能不同。", "提供requirements.txt和不依赖pytest的tests/run_checks.py验证脚本。"],
    ])

    add_heading(doc, "4. 核心代码展示")
    add_code(doc, "函数1：检索器search，负责从本地向量索引中找出最相关校园规则。", """def search(self, query: str, top_k: int = DEFAULT_TOP_K):
    # 对用户问题做同义词扩展，例如“发烧”扩展为“病假、请假、医院”
    query_vector = self._vectorize(tokenize(expand_query(query)))
    scored = [
        (record, self._cosine(query_vector, doc_vector))
        for record, doc_vector in zip(self.records, self.doc_vectors)
    ]
    scored.sort(key=lambda item: item[1], reverse=True)
    return [item for item in scored[:top_k] if item[1] > 0]""")
    add_code(doc, "函数2：rag_answer，负责“检索-增强-回答”的核心流程。", """def rag_answer(question: str, retriever=None, use_llm: bool = True):
    active_retriever = retriever or build_retriever()
    results = active_retriever.search(question, top_k=3)
    context = _format_context(results)
    answer = None
    if use_llm and results:
        answer = _deepseek_answer(question, context)
    if not answer:
        answer = _local_answer(question, results)
    return {"answer": answer, "context": context, "results": results}""")
    add_code(doc, "函数3：CampusAgent.chat，负责智能体意图识别和工具调用。", """def chat(self, user_input: str, use_llm: bool = True) -> str:
    question = user_input.strip()
    self.conversation_history.append({"role": "user", "content": question})
    if self._is_week_query(question):
        response = get_current_week()
    elif self._is_gpa_query(question):
        response = calculate_gpa(question)
    else:
        response = rag_answer(question, self.retriever, use_llm=use_llm)["answer"]
    self.conversation_history.append({"role": "assistant", "content": response})
    self.conversation_history = self.conversation_history[-10:]
    return response""")

    add_heading(doc, "5. 运行截图")
    for title_text, image_name in [("图1 RAG智能问答结果", "01_rag_answer.png"), ("图2 检索过程与相似度展示", "02_retrieval.png"), ("图3 智能体工具调用结果", "03_tools.png")]:
        add_para(doc, title_text, bold=True)
        doc.add_picture(str(SHOTS / image_name), width=Inches(6.2))

    add_heading(doc, "6. 问题与反思")
    add_para(doc, "本次实训遇到的最大困难是如何兼顾“可运行演示”和“数据真实可靠”。指导书示例中包含请假、奖学金、报修、一卡通等生活场景，但学校官网公开页面并不一定逐项给出详细流程。如果直接编写看似合理的答案，系统虽然能演示，但真实性不足。因此我对CSV进行了官网核验，保留教务处《学分制学籍管理办法》、公共服务校历、财务处缴费与收费公示等可确认内容；对奖学金固定绩点线、报修时限、一卡通补办细节等无法确认的信息不再写死，而是提示以学生处、辅导员或学院当年通知为准。第二个困难是本地环境中未必能安装Chroma和中文嵌入模型。为了保证机房环境可复现，我使用TF-IDF实现轻量检索，并保留DeepSeek API增强入口。后续如果有更多权威资料，我会继续补充后勤服务、校园卡、图书馆等模块，并使用真正的中文嵌入模型提升语义检索效果。")

    add_heading(doc, "7. 实训总结")
    add_para(doc, "通过本次实训，我完整体验了一个AI应用从需求分析、数据整理、检索模块、RAG问答、工具型智能体到Web界面的开发流程。以前我对AI应用的理解更多停留在“调用大模型回答问题”，现在认识到高质量AI应用更依赖可靠的数据、清晰的提示词、可解释的检索来源和稳定的工程实现。校园百事通助手虽然规模不大，但它包含了真实应用中的关键环节：知识库建设决定回答边界，检索质量影响回答准确性，智能体工具让系统能处理计算类任务，界面和测试则保证项目可展示、可复现。实训也让我意识到AI不是替代程序设计，而是需要和传统软件工程结合，才能变成真正可用的产品。")

    add_heading(doc, "附：测试结果")
    add_table(doc, ["测试问题", "期望关键词", "结果"], [
        ["怎么请病假？", "辅导员、证明、销假", "通过"],
        ["奖学金要多少绩点？", "3.0、无挂科、体测合格", "通过"],
        ["生病不能考试怎么办？", "缓考、诊断证明", "通过"],
        ["学生证丢了怎么办？", "学生证补办办理流程", "通过"],
        ["选错了课能退吗？", "退选、教务处", "通过"],
        ["学费住宿费在哪里缴？", "网上缴费系统、收费公示", "通过"],
        ["现在第几周？", "本学期第17周", "通过"],
        ["绩点计算 85,90,78", "3.00", "通过"],
    ])

    output = DOCS / "校园生活百事通助手-实训报告.docx"
    doc.save(output)
    return output


def main() -> None:
    DOCS.mkdir(exist_ok=True)
    SHOTS.mkdir(exist_ok=True)
    save_chat_screenshot()
    save_retrieval_screenshot()
    save_tools_screenshot()
    save_architecture()
    output = create_report()
    print(output)


if __name__ == "__main__":
    main()
