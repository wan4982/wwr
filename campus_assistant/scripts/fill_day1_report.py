from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


REPORT = Path(r"C:\Users\wan\Desktop\202431243 王婉如 实训报告.docx")
FALLBACK_REPORT = Path(r"D:\StudyProjects\campus_assistant\docs\202431243 王婉如 实训报告-第1天已填.docx")


def set_cell_text(cell, text: str, font_size: float = 10.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.line_spacing = 1.2
    paragraph.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(text.split("\n")):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(font_size)


def main() -> None:
    doc = Document(REPORT)
    table = doc.tables[0]

    set_cell_text(table.cell(0, 1), "2026年6月22日")
    set_cell_text(table.cell(0, 3), "机房 / 宿舍")

    set_cell_text(
        table.cell(1, 1),
        "完成校园生活百事通助手第1天数据收集与整理：确定项目主题，梳理请假、奖学金、报修、一卡通、选课、宿舍等校园生活场景，整理形成 data/campus_data.csv 知识库，共48条问答记录。",
    )

    set_cell_text(table.cell(2, 0), "一、遇到的问题及其解决方法")
    set_cell_text(table.cell(3, 0), "问题描述")
    set_cell_text(table.cell(3, 2), "解决方法")

    problems = [
        (
            "校园规则来源较多，官网、操作手册和管理规定分散，直接收集后格式不统一。",
            "统一整理为 id、category、question、answer、source 五个字段，便于后续检索和生成回答。",
        ),
        (
            "同一类问题存在多种问法，例如“怎么请病假”“怎么请假”“请假结束后做什么”。",
            "为常见场景设计标准问题，并在答案中保留关键步骤，如今日校园申请、材料上传、销假等。",
        ),
        (
            "部分内容如果只写一句话，后续问答容易不完整。",
            "将请假类数据拆成病假、事假、销假、特殊情况等多条问答，后续由程序组合成流程型回答。",
        ),
        (
            "CSV 中包含中文标点、引号和 APP 路径，容易导致读取格式出错。",
            "使用标准 CSV 格式保存，对包含引号的内容进行转义，并用程序读取检查48条数据是否完整。",
        ),
        (
            "数据量虽然不大，但需要考虑后续向量检索效果。",
            "每条答案保持短小清晰，把一条校园规则作为一个知识片段，暂时不需要再进行文件切分。",
        ),
    ]
    for row_index, (problem, solution) in enumerate(problems, start=4):
        set_cell_text(table.cell(row_index, 0), problem)
        set_cell_text(table.cell(row_index, 2), solution)

    set_cell_text(table.cell(9, 0), "二、关键实现代码")
    set_cell_text(
        table.cell(10, 0),
        """# data/campus_data.csv 的字段结构
id, category, question, answer, source

# src/data_loader.py：读取校园知识库
def load_records(path=DATA_FILE):
    records = []
    with path.open("r", encoding="utf-8-sig", newline="") as file:
        for row in csv.DictReader(file):
            records.append(CampusRecord(
                id=int(row["id"]),
                category=row["category"].strip(),
                question=row["question"].strip(),
                answer=row["answer"].strip(),
                source=row["source"].strip(),
            ))
    return records""",
        font_size=9,
    )

    set_cell_text(table.cell(11, 0), "三、实训小结（不少于100字）")
    set_cell_text(
        table.cell(12, 0),
        "第一天主要完成了项目选题和数据准备工作。我选择“校园生活百事通助手”作为实训项目，目标是让系统能够回答学生在校内生活中经常遇到的问题。当天重点不是直接写界面，而是先整理知识库，因为RAG问答的准确性很大程度取决于数据质量。我根据校园生活场景收集并整理了请假、奖学金、报修、一卡通、选课、宿舍管理等内容，把资料统一转换为CSV格式，并保留来源字段。通过这一步，我理解了RAG中“检索增强”的基础：只有把知识片段整理清楚，后续向量检索和智能问答才有可靠依据。",
    )

    try:
        doc.save(REPORT)
        print(f"已填写第1天实训记录：{REPORT}")
    except PermissionError:
        FALLBACK_REPORT.parent.mkdir(parents=True, exist_ok=True)
        doc.save(FALLBACK_REPORT)
        print(f"原文件被占用，已另存为：{FALLBACK_REPORT}")


if __name__ == "__main__":
    main()
